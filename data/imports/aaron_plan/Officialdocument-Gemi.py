# -*- coding: utf-8 -*-
# Officialdocument-Gemi.py

# ==============================================================================
# 主要功能與修正歷程 V5 (Approx. 2025-04-12)
# ==============================================================================
# 1.  資料來源: 從指定資料夾讀取多種格式公文清單 CSV 檔案。
#     - 支持檔名格式: YYYY-MM-DD_ (+ receiveList/sendList/preceiveList/psendList) + .csv
#     - 自動判斷文件類型 (收/發文) 與形式 (電子/紙本)。
#     - 兼容 UTF-8 及 Big5 編碼。
# 2.  數據處理:
#     - 清理欄位值 (去除空白)。
#     - 民國日期 ('公文日期') 轉換為西元日期 ('日期')。
#     - 根據設定檔映射來源欄位 ('使用者確認'/'收文日期' -> '收文日期'; '狀態' -> '收發狀態')。
#     - 處理跨日期重複數據：合併所有數據後，保留每個唯一記錄 (以 '公文字號'+'文件類型'+'受文單位' 識別) 的最新版本 (基於 '系統輸出日期')。
# 3.  流水號生成:
#     - 為最終唯一記錄集生成從 1 開始的連續 '流水號'。
#     - 流水號基於公文實際日期 ('日期' 欄位) 升序排列。
# 4.  Notion 同步 (Upsert):
#     - 使用 '公文字號'+'文件類型'+'受文單位' 作為唯一鍵檢查 Notion。
#     - 創建 Notion 中不存在的新記錄。
#     - 更新 Notion 中已存在且內容有變化的記錄 (包括流水號)。
#     - 跳過數據完全相同的記錄。
# 5.  Notion 清理 (歸檔):
#     - 自動歸檔 Notion 中存在但本地最新數據中已不存在的記錄。
#     - **依賴 Notion 資料庫中必須存在名為 `archived` (或 config 中指定名稱) 的 `Checkbox` 屬性。**
# 6.  健壯性與配置:
#     - 使用 .env 檔案配置 Notion Token 和 Database ID。
#     - 使用 config.json 外部檔案配置檔案路徑、欄位映射、Notion 屬性名、日誌級別等。
#     - 包含 API 調用自動重試機制 (tenacity)。
#     - 提供詳細日誌輸出到 console 及 document_processing.log 文件。
# 7.  用戶體驗:
#     - 在處理文件、同步、歸檔等步驟顯示進度條 (tqdm)。
#     - 可選功能：自動生成 Word 格式的操作說明文件 (僅在文件不存在時)。
# ==============================================================================


import pandas as pd
import os
from notion_client import Client, APIResponseError
from tenacity import retry, stop_after_attempt, wait_exponential
import logging
from typing import Dict, Any, Optional, List
from datetime import datetime
import time
from dotenv import load_dotenv
import json
from tqdm import tqdm
try:
    import docx # ---!!! 需要安裝: pip install python-docx !!!---
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# --- Setup Section ---
load_dotenv()
logging.basicConfig(
    level=logging.INFO, # DEBUG for detailed logs
    format='%(asctime)s - %(levelname)s - [%(funcName)s:%(lineno)d] %(message)s',
    handlers=[
        logging.FileHandler("document_processing.log", encoding='utf-8', mode='w'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
logging.getLogger("notion_client").setLevel(logging.WARNING)
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("tenacity").setLevel(logging.WARNING)

# --- Load Configuration ---
CONFIG_FILE = "config.json"
DEFAULT_CONFIG = {
    "folder_path": "data",
    "log_level": "INFO",
    "rate_limit_delay": 0.4,
    "expected_source_columns": ['公文日期', '類別', '字', '文號', '主旨', '發文單位', '受文單位', '狀態', '收文日期', '使用者確認', '備註', '承攬案件'],
    "field_config": {
        'sent_electronic': {'receive_date_field': '使用者確認', 'status_field': '狀態'},
        'receive_electronic': {'receive_date_field': '收文日期', 'status_field': '狀態'},
        'sent_paper': {'receive_date_field': '使用者確認', 'status_field': '狀態'},
        'receive_paper': {'receive_date_field': '收文日期', 'status_field': '狀態'},
    },
    "final_columns": ['流水號', '文件類型', '公文字號', '日期', '公文日期', '類別', '字', '文號', '主旨', '發文單位', '受文單位', '收發狀態', '收文日期', '發文形式', '備註', '承攬案件', '系統輸出日期'],
    "notion_property_names": {
        "流水號": "流水號", "文件類型": "文件類型", "公文字號": "公文字號", "日期": "日期",
        "公文日期": "公文日期", "類別": "類別", "字": "字", "文號": "文號", "主旨": "主旨",
        "發文單位": "發文單位", "受文單位": "受文單位", "收發狀態": "收發狀態", "收文日期": "收文日期",
        "發文形式": "發文形式", "備註": "備註", "承攬案件": "承攬案件", "系統輸出日期": "系統輸出日期",
        "archived": "archived"
    }
}
try:
    with open(CONFIG_FILE, 'r', encoding='utf-8') as f: config = json.load(f)
    for key, value in DEFAULT_CONFIG.items(): config.setdefault(key, value)
    print(f"成功從 {CONFIG_FILE} 載入設定。")
except FileNotFoundError: print(f"警告：未找到設定檔 {CONFIG_FILE}，使用預設設定。"); config = DEFAULT_CONFIG
except json.JSONDecodeError: print(f"錯誤：設定檔 {CONFIG_FILE} 格式錯誤。使用預設設定。"); config = DEFAULT_CONFIG
except Exception as e: print(f"讀取設定檔時發生錯誤: {e}，使用預設設定。"); config = DEFAULT_CONFIG

log_level_str = config.get("log_level", "INFO").upper(); log_level = getattr(logging, log_level_str, logging.INFO)
logging.getLogger().setLevel(log_level)


class NotionDocumentSync:
    def __init__(self, config: Dict[str, Any]):
        """初始化"""
        try:
            self.config = config
            self.folder_path = self.config['folder_path']
            self.expected_source_columns = self.config['expected_source_columns']
            self.field_config = self.config['field_config']
            self.rate_limit_delay = self.config['rate_limit_delay']
            self.final_columns = self.config['final_columns']
            self.notion_prop_names = self.config['notion_property_names']

            if not os.path.isdir(self.folder_path): raise FileNotFoundError(f"指定的資料夾路徑不存在: {self.folder_path}")
            logger.info(f"設定資料夾路徑為: {self.folder_path}")
            self.setup_notion()
            logger.info("公文同步處理器初始化完成")
        except Exception as e: logger.critical(f"初始化失敗: {e}", exc_info=True); raise

    def setup_notion(self):
        """設定 Notion 連接"""
        self.notion_token = os.getenv("NOTION_TOKEN"); self.database_id = os.getenv("NOTION_DATABASE_ID")
        if not self.notion_token: raise ValueError("錯誤：未找到 NOTION_TOKEN")
        if not self.database_id: raise ValueError("錯誤：未找到 NOTION_DATABASE_ID")
        logger.debug(f"讀取到 Token: ...{self.notion_token[-4:]}"); logger.debug(f"讀取到 Database ID: {self.database_id}")
        if not self.validate_notion_id(self.database_id): raise ValueError(f"無效的資料庫 ID 格式: {self.database_id}")
        try:
            self.notion = Client(auth=self.notion_token, notion_version="2022-06-28", log_level=logging.WARNING)
            logger.info("Notion 客戶端初始化成功")
            db_info = self.notion.databases.retrieve(database_id=self.database_id)
            title_prop_name = self.notion_prop_names.get('公文字號', '公文字號')
            db_title = "未知資料庫標題"
            if db_info and db_info.get('title'):
                 try: db_title = db_info['title'][0].get('plain_text', '無標題')
                 except: pass
            logger.info(f"成功連接至 Notion 資料庫: '{db_title}' (ID: ...{self.database_id[-4:]})")
        except APIResponseError as api_error:
             error_body = {};
             try: error_body = json.loads(api_error.body)
             except: error_body = {"raw_body": str(api_error.body)}
             logger.critical(f"Notion API 錯誤導致初始化失敗: Status={api_error.status}, Code={api_error.code}, Body={error_body}", exc_info=False)
             raise
        except Exception as e: logger.critical(f"Notion 初始化失敗: {e}", exc_info=True); raise

    def validate_notion_id(self, notion_id: str) -> bool:
        """驗證 Notion ID 格式"""
        if not notion_id: return False
        notion_id_cleaned = notion_id.replace('-', ''); return len(notion_id_cleaned) == 32 and all(c in '0123456789abcdefABCDEF' for c in notion_id_cleaned)

    def get_date_files(self) -> Dict[str, List[Dict[str, str]]]:
        """掃描資料夾"""
        date_files: Dict[str, List[Dict[str, str]]] = {}
        file_patterns = {"_receivelist.csv": "receive_electronic", "_sendlist.csv": "sent_electronic", "_preceivelist.csv": "receive_paper", "_psendlist.csv": "sent_paper"}
        if not os.path.isdir(self.folder_path): logger.error(f"指定的資料夾路徑不存在: {self.folder_path}"); return {}
        try:
            all_items = os.listdir(self.folder_path)
            logger.info(f"掃描資料夾 '{self.folder_path}' 中的 {len(all_items)} 個項目...")
            for filename in tqdm(all_items, desc="掃描檔案", unit="項目", leave=False):
                fn_lower = filename.lower(); matched_type = None
                for pattern, type_name in file_patterns.items():
                    if fn_lower.endswith(pattern): matched_type = type_name; break
                if matched_type:
                    try:
                        date_str = filename[:10]; datetime.strptime(date_str, '%Y-%m-%d')
                        file_path = os.path.join(self.folder_path, filename)
                        if os.path.isfile(file_path):
                            if date_str not in date_files: date_files[date_str] = []
                            date_files[date_str].append({"path": file_path, "type": matched_type}); logger.debug(f"找到檔案: {filename} (類型: {matched_type})")
                    except (ValueError, IndexError): logger.warning(f"檔名格式錯誤: {filename}")
                    except Exception as e: logger.error(f"處理檔名 '{filename}' 出錯: {e}")
        except Exception as e: logger.error(f"掃描 '{self.folder_path}' 出錯: {e}", exc_info=True); return {}
        total_file_count = sum(len(f) for f in date_files.values())
        logger.info(f"在 '{self.folder_path}' 找到 {len(date_files)} 個日期的共 {total_file_count} 個公文檔案。")
        return date_files

    def safe_get(self, row: pd.Series, column: str, default: Any = '') -> Any:
        """安全獲取並清理字串"""
        try:
            if column not in row.index: return default
            value = row[column];
            if pd.isna(value): return default
            return value.strip() if isinstance(value, str) else value
        except Exception: return default

    def convert_roc_date(self, roc_date: Any, output_format: str = '%Y-%m-%d') -> str:
        """轉換民國日期"""
        try:
            if pd.isna(roc_date): return ''
            date_str = str(roc_date).strip().replace("中華民國", "").replace("民國", "").strip()
            if not all(x in date_str for x in ['年', '月', '日']): return ''
            parts = date_str.split('年');
            if len(parts) < 2: return ''
            year_str, month_day_str = parts[0], parts[1]; month_parts = month_day_str.split('月');
            if len(month_parts) < 2: return ''
            month_str, day_parts_str = month_parts[0], month_parts[1]; day_parts = day_parts_str.split('日');
            if not day_parts: return ''
            day_str = day_parts[0]
            if not (year_str.isdigit() and month_str.isdigit() and day_str.isdigit()): return ''
            year, month, day = int(year_str), int(month_str), int(day_str)
            if not (1 <= month <= 12 and 1 <= day <= 31): return ''
            ad_year = year + 1911; dt = datetime(ad_year, month, day)
            return dt.strftime(output_format)
        except (ValueError, TypeError): return ''
        except Exception as e: logger.error(f"日期轉換錯誤: {roc_date} - {e}", exc_info=False); return ''

    def load_csv_data(self, filepath: str, doc_type_detail: str) -> pd.DataFrame:
        """載入 CSV 檔案"""
        df = pd.DataFrame(); encodings_to_try = ['utf-8-sig', 'big5']; loaded = False; last_error = None
        for encoding in encodings_to_try:
            try:
                df = pd.read_csv(filepath, skiprows=3, header=0, encoding=encoding, dtype=str, skipinitialspace=True)
                df.dropna(how='all', inplace=True); df.columns = df.columns.str.strip(); df.fillna('', inplace=True)
                logger.debug(f"成功使用 '{encoding}' 載入 {doc_type_detail} 清單: {os.path.basename(filepath)}")
                loaded = True; break
            except UnicodeDecodeError: last_error = f"UnicodeDecodeError with {encoding}"; logger.debug(f"編碼錯誤 ({encoding})，嘗試下一個: {os.path.basename(filepath)}"); continue
            except FileNotFoundError: logger.error(f"{doc_type_detail} 清單檔案未找到: {filepath}"); return pd.DataFrame()
            except Exception as e: logger.error(f"使用 '{encoding}' 載入 {doc_type_detail} 時出錯: {e}", exc_info=True); last_error = e
        if not loaded: logger.error(f"嘗試所有編碼後仍無法載入檔案: {os.path.basename(filepath)}. 最後錯誤: {last_error}"); return pd.DataFrame()
        loaded_columns = df.columns.tolist(); missing_cols = [col for col in self.expected_source_columns if col not in loaded_columns]
        if missing_cols: logger.warning(f"CSV '{os.path.basename(filepath)}' ({doc_type_detail}) 可能缺少來源欄位: {', '.join(missing_cols)}。")
        logger.debug(f"載入 {doc_type_detail} 清單完成: {os.path.basename(filepath)}，共 {len(df)} 筆資料")
        return df

    def prepare_data(self, df: pd.DataFrame, doc_type_detail: str, date_str: str) -> pd.DataFrame:
        """準備公文資料"""
        final_columns = self.final_columns
        try:
            if df.empty: logger.warning(f"傳入 prepare_data 的 DataFrame 為空 ({doc_type_detail}, {date_str})。"); return pd.DataFrame(columns=final_columns).fillna('')
            df_prep = df.copy()
            base_doc_type = '發文' if 'sent' in doc_type_detail else '收文'
            df_prep['文件類型'] = str(base_doc_type); df_prep['發文形式'] = str('紙本' if 'paper' in doc_type_detail else '電子'); df_prep['系統輸出日期'] = str(date_str); df_prep['流水號'] = ''
            for col in ['類別', '字', '文號', '主旨', '發文單位', '備註', '承攬案件', '公文日期']: df_prep[col] = df_prep.apply(lambda row: str(self.safe_get(row, col)), axis=1)

            # --- !!! 修改點：統一使用 "...字第...號" 格式 !!! ---
            df_prep['公文字號'] = df_prep.apply(lambda row: f"{row['字']}字第{row['文號']}號" if row.get('字') and row.get('文號') else '', axis=1).astype(str)
            # ---------------------------------------------

            df_prep['日期'] = df_prep.apply(lambda row: self.convert_roc_date(row['公文日期']), axis=1).astype(str)
            config_detail = self.field_config.get(doc_type_detail, {}); status_field = config_detail.get('status_field', '狀態'); receive_date_field = config_detail.get('receive_date_field', '收文日期')
            df_prep['收發狀態'] = df_prep.apply(lambda row: str(self.safe_get(row, status_field)), axis=1)
            df_prep['收文日期'] = df_prep.apply(lambda row: str(self.safe_get(row, receive_date_field)), axis=1)
            df_prep['受文單位'] = df_prep.apply(lambda row: str(self.safe_get(row, '受文單位')), axis=1)
            df_prep['受文單位'] = df_prep['受文單位'].str.replace(r'\s*\n\s*', ' ', regex=True).str.replace(r'\s+', ' ', regex=True).str.strip()
            df_prep['受文單位'] = df_prep['受文單位'].str.replace(r'(\S)\(', r'\1 (', regex=True)
            for col in final_columns:
                if col not in df_prep.columns: df_prep[col] = ''
            df_prep = df_prep.fillna('')
            for col in final_columns: df_prep[col] = df_prep[col].astype(str)
            return df_prep[final_columns]
        except Exception as e:
            logger.error(f"準備 {doc_type_detail} 資料 ({date_str}) 時出錯: {e}", exc_info=True)
            return pd.DataFrame(columns=final_columns).fillna('')

    # --- Helper Methods for Notion ---
    def _add_select_prop(self, properties: Dict[str, Any], prop_name_internal: str, value: Any):
        notion_name = self.notion_prop_names.get(prop_name_internal)
        if not notion_name: logger.warning(f"Config未定義Notion屬性名: {prop_name_internal}"); return
        str_value = str(value).strip() if value is not None else '';
        if str_value: properties[notion_name] = {"select": {"name": str_value}}

    def _add_rich_text_prop(self, properties: Dict[str, Any], prop_name_internal: str, value: Any):
        notion_name = self.notion_prop_names.get(prop_name_internal)
        if not notion_name: logger.warning(f"Config未定義Notion屬性名: {prop_name_internal}"); return
        str_value = str(value).strip() if value is not None else '';
        if str_value: properties[notion_name] = {"rich_text": [{"text": {"content": str_value}}]}

    def _get_clear_payload(self, prop_type: str) -> Optional[Dict | None]:
        """根據屬性類型返回用於清空該屬性的 Notion API payload"""
        if prop_type == 'select': return None # Select 用 None 清除
        elif prop_type == 'date': return {'date': None}
        elif prop_type == 'number': return {'number': None}
        elif prop_type == 'rich_text': return {'rich_text': []}
        elif prop_type == 'multi_select': return {'multi_select': []}
        elif prop_type == 'relation': return {'relation': []}
        elif prop_type == 'people': return {'people': []}
        return None # 對於不支持清空或未知的類型，返回 None

    # --- Notion API Interaction ---
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1.5, min=4, max=10), reraise=True)
    def _notion_api_call(self, api_func, **kwargs):
        time.sleep(self.rate_limit_delay)
        logger.debug(f"執行 Notion API: {api_func.__name__} (PageID: {kwargs.get('page_id', 'N/A')})")
        try: response = api_func(**kwargs); return response
        except APIResponseError as e:
            error_body_str = "Unknown body";
            try: error_body_str = e.body.decode('utf-8');
            except: pass;
            logger.error(f"Notion API 錯誤 ({api_func.__name__}): Status={e.status}, Code={e.code}, Body={error_body_str}")
            raise

    def create_notion_page(self, properties: Dict[str, Any]) -> Optional[str]:
        title_prop_name = self.notion_prop_names.get('公文字號', '公文字號')
        doc_id_content = properties.get(title_prop_name, {}).get('title', [{}])[0].get('text', {}).get('content', '未知') if properties.get(title_prop_name) else '未知'
        try:
            response = self._notion_api_call(self.notion.pages.create, parent={"database_id": self.database_id}, properties=properties)
            page_id = response.get('id'); logger.debug(f"建立頁面回應 ({doc_id_content}): ID={page_id}")
            if page_id: logger.info(f"成功建立頁面: {doc_id_content} (ID: {page_id})"); return page_id
            else: logger.error(f"建立頁面失敗，API未返回頁面ID: {doc_id_content}"); return None
        except Exception as e: logger.error(f"嘗試建立頁面最終失敗 ({doc_id_content})，原因: {str(e)}"); return None

    def update_notion_page(self, page_id: str, properties: Dict[str, Any]) -> bool:
        try:
            self._notion_api_call(self.notion.pages.update, page_id=page_id, properties=properties)
            logger.info(f"成功更新頁面ID: {page_id}"); return True
        except Exception as e: logger.error(f"嘗試更新頁面最終失敗 (ID: {page_id})，原因: {str(e)}"); return False

    def archive_notion_page(self, page_id: str) -> bool:
        """Archive an existing page in Notion."""
        try:
            self._notion_api_call(self.notion.pages.update, page_id=page_id, archived=True) # Directly set archived=True
            archive_prop_name = self.notion_prop_names.get('archived', 'archived') # Name used only for logging if needed
            logger.info(f"成功歸檔頁面ID: {page_id}")
            return True
        except APIResponseError as e:
             error_body_str = str(e.body)
             archive_prop_name = self.notion_prop_names.get('archived', 'archived')
             if archive_prop_name in error_body_str and "Could not find property" in error_body_str: logger.error(f"歸檔失敗 (ID: {page_id}): Notion 資料庫缺少名為 '{archive_prop_name}' 的 Checkbox 屬性。", exc_info=False)
             else: logger.error(f"嘗試歸檔頁面時發生 API 錯誤 (ID: {page_id})，原因: Status={e.status}, Code={e.code}, Body={error_body_str}")
             return False
        except Exception as e: logger.error(f"嘗試歸檔頁面時發生未知錯誤 (ID: {page_id})，原因: {str(e)}"); return False

    def _parse_flexible_date(self, date_str: Any) -> Optional[str]:
        """解析日期為YYYY-MM-DD"""
        if pd.isna(date_str) or date_str is None: return None
        date_str = str(date_str).strip();
        if not date_str: return None
        formats_to_try = ['%Y-%m-%d %H:%M:%S', '%Y/%m/%d %H:%M:%S', '%Y-%m-%d %H:%M', '%Y/%m/%d %H:%M', '%Y-%m-%d', '%Y/%m/%d']
        for fmt in formats_to_try:
            try: dt = datetime.strptime(date_str, fmt); return dt.strftime('%Y-%m-%d')
            except ValueError: continue
        try: # ISO
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00')); return dt.strftime('%Y-%m-%d')
        except ValueError: logger.debug(f"無法解析日期格式: '{date_str}'"); return None

    def build_notion_properties(self, row: pd.Series) -> Dict[str, Any]:
        """構建 Notion 屬性字典"""
        properties = {}; title_prop_name = self.notion_prop_names.get('公文字號', '公文字號'); doc_id_for_log = self.safe_get(row, '公文字號', '未知公文字號')
        try:
            doc_id = self.safe_get(row, '公文字號'); # Contains ...字第... format
            if not doc_id: logger.warning(f"公文字號為空..."); return {}
            properties[title_prop_name] = {"title": [{"text": {"content": doc_id}}]}

            self._add_select_prop(properties, "文件類型", self.safe_get(row, '文件類型')); self._add_select_prop(properties, "類別", self.safe_get(row, '類別')); self._add_select_prop(properties, "發文單位", self.safe_get(row, '發文單位'))
            self._add_select_prop(properties, "受文單位", self.safe_get(row, '受文單位')); self._add_select_prop(properties, "收發狀態", self.safe_get(row, '收發狀態')); self._add_select_prop(properties, "發文形式", self.safe_get(row, '發文形式'))
            self._add_select_prop(properties, "承攬案件", self.safe_get(row, '承攬案件'))

            serial_no_prop_name = self.notion_prop_names.get('流水號', '流水號')
            serial_no_str = self.safe_get(row, '流水號')
            if serial_no_str:
                 try: properties[serial_no_prop_name] = {"number": int(serial_no_str)}; logger.debug(f"Adding Number '{serial_no_prop_name}': {serial_no_str}")
                 except (ValueError, TypeError): logger.warning(f"'{serial_no_prop_name}' ('{serial_no_str}') 不是有效數字 for {doc_id_for_log}")

            date_prop_name = self.notion_prop_names.get('日期', '日期')
            date_val = self.safe_get(row, '日期');
            if date_val: properties[date_prop_name] = {"date": {"start": date_val}}

            receive_date_prop_name = self.notion_prop_names.get('收文日期', '收文日期')
            receive_date_raw = self.safe_get(row, '收文日期'); receive_date_iso = self._parse_flexible_date(receive_date_raw)
            if receive_date_iso: properties[receive_date_prop_name] = {"date": {"start": receive_date_iso}}

            output_date_prop_name = self.notion_prop_names.get('系統輸出日期', '系統輸出日期')
            output_date_val = self.safe_get(row, '系統輸出日期')
            if output_date_val:
                try: datetime.strptime(output_date_val, '%Y-%m-%d'); properties[output_date_prop_name] = {"date": {"start": output_date_val}}; logger.debug(f"Adding Date '{output_date_prop_name}': {output_date_val}")
                except (ValueError, TypeError): logger.warning(f"'{output_date_prop_name}' ('{output_date_val}') 不是有效YYYY-MM-DD 格式 for {doc_id_for_log}")

            self._add_rich_text_prop(properties, "公文日期", self.safe_get(row, '公文日期')); self._add_rich_text_prop(properties, "字", self.safe_get(row, '字')); self._add_rich_text_prop(properties, "文號", self.safe_get(row, '文號'))
            self._add_rich_text_prop(properties, "主旨", self.safe_get(row, '主旨')); self._add_rich_text_prop(properties, "備註", self.safe_get(row, '備註'))

        except Exception as e: logger.error(f"構建 Notion 屬性時出錯 ({doc_id_for_log}): {e}", exc_info=True); return {}
        logger.debug(f"最終構建屬性 ({doc_id_for_log}): {list(properties.keys())}")
        return properties

    def needs_update(self, existing_page: Dict[str, Any], new_properties: Dict[str, Any]) -> bool:
        """檢查現有頁面是否需要更新"""
        try:
            existing_props = existing_page.get('properties', {})
            title_prop_name = self.notion_prop_names.get('公文字號', '公文字號')
            doc_id_for_log = "未知頁面";
            if title_prop_name in existing_props and existing_props[title_prop_name].get('title'):
                try: doc_id_for_log = existing_props[title_prop_name]['title'][0].get('text', {}).get('content', '未知')
                except: pass
            needs_update_flag = False; update_reasons = []
            for internal_key, new_value_dict in new_properties.items():
                notion_key = self.notion_prop_names.get(internal_key);
                if not notion_key: continue
                if notion_key not in existing_props:
                     prop_type = list(new_value_dict.keys())[0]; new_val = None
                     if prop_type == "select": new_val = new_value_dict.get('select', {}).get('name')
                     elif prop_type == "date": new_val = new_value_dict.get('date', {}).get('start'); new_val = new_val[:10] if new_val else ''
                     elif prop_type == "number": new_val = new_value_dict.get('number')
                     elif prop_type == "rich_text": new_list = new_value_dict.get('rich_text', []); new_val = new_list[0].get('text', {}).get('content') if new_list else ''
                     if new_val is not None and str(new_val).strip() != '': update_reasons.append(f"新增欄位 '{notion_key}'"); needs_update_flag = True
                     continue
                existing_prop_dict = existing_props[notion_key]; prop_type = existing_prop_dict.get('type')
                existing_val_str, new_val_str = '', ''
                try:
                    existing_val, new_val = None, None
                    if prop_type == "title": existing_val = existing_prop_dict.get('title', [{}])[0].get('text', {}).get('content'); new_val = new_value_dict.get('title', [{}])[0].get('text', {}).get('content')
                    elif prop_type == "rich_text": existing_list = existing_prop_dict.get('rich_text', []); existing_val = existing_list[0].get('text', {}).get('content') if existing_list else ''; new_list = new_value_dict.get('rich_text', []); new_val = new_list[0].get('text', {}).get('content') if new_list else ''
                    elif prop_type == "select": existing_val = existing_prop_dict.get('select', {}).get('name'); new_val = new_value_dict.get('select', {}).get('name')
                    elif prop_type == "date": existing_val = existing_prop_dict.get('date', {}).get('start'); new_val = new_value_dict.get('date', {}).get('start'); existing_val = existing_val[:10] if existing_val else ''; new_val = new_val[:10] if new_val else ''
                    elif prop_type == "number":
                         existing_val = existing_prop_dict.get('number'); new_val = new_value_dict.get('number')
                         try: existing_float = float(existing_val) if existing_val is not None else None; new_float = float(new_val) if new_val is not None else None
                         except (ValueError, TypeError): existing_float, new_float = object(), object()
                         if existing_float != new_float: update_reasons.append(f"欄位 '{notion_key}' 不同: '{existing_val}' -> '{new_val}'"); needs_update_flag = True
                         continue
                    else: continue
                    existing_val_str = str(existing_val or ''); new_val_str = str(new_val or '')
                    if existing_val_str != new_val_str: update_reasons.append(f"欄位 '{notion_key}' 不同: '{existing_val_str}' -> '{new_val_str}'"); needs_update_flag = True
                except (KeyError, IndexError, TypeError) as comp_err: logger.warning(f"{doc_id_for_log}: 比較 {notion_key} ({prop_type}) 出錯: {comp_err}. 假定需更新."); update_reasons.append(f"比較欄位 '{notion_key}' 出錯"); needs_update_flag = True
            new_prop_notion_names = set(self.notion_prop_names.values())
            for key in existing_props.keys():
                if key not in new_prop_notion_names and key != title_prop_name :
                     existing_prop_dict = existing_props[key]; prop_type = existing_prop_dict.get('type'); has_value = False
                     if prop_type == "select" and existing_prop_dict.get('select'): has_value = True
                     elif prop_type == "date" and existing_prop_dict.get('date'): has_value = True
                     elif prop_type == "rich_text" and existing_prop_dict.get('rich_text'): has_value = True
                     elif prop_type == "number" and existing_prop_dict.get('number') is not None: has_value = True
                     if has_value: update_reasons.append(f"欄位 '{key}' 需要清空"); needs_update_flag = True
            if needs_update_flag: logger.info(f"頁面 {doc_id_for_log}: 需要更新，原因: {'; '.join(update_reasons)}")
            else: logger.debug(f"頁面 {doc_id_for_log}: 無需更新。")
            return needs_update_flag
        except Exception as e: logger.error(f"檢查更新需求時出錯 ({doc_id_for_log}): {e}", exc_info=True); return True

    def connect_notion(self) -> bool:
        """連線到 Notion 資料庫並驗證"""
        try: self.notion.databases.retrieve(database_id=self.database_id); logger.info("Notion 連線驗證成功。"); return True
        except APIResponseError as e: logger.error(f"Notion API 錯誤導致連線驗證失敗: Status={e.status}, Code={e.code}, Body={e.body}"); return False
        except Exception as e: logger.error(f"Notion 連線驗證失敗: {e}", exc_info=True); return False

    def get_existing_documents(self) -> Dict[str, Dict[str, Any]]:
        """從 Notion 取得現有(未歸檔)公文資料"""
        existing_docs: Dict[str, Dict[str, Any]] = {}; next_cursor: Optional[str] = None; page_count = 0; total_retrieved = 0
        logger.info("開始從 Notion 獲取現有文件(僅限未歸檔)...")
        use_filter = True
        archived_prop_name = self.notion_prop_names.get('archived', 'archived'); title_prop_name = self.notion_prop_names.get('公文字號', '公文字號'); doc_type_prop_name = self.notion_prop_names.get('文件類型', '文件類型'); recipient_prop_name = self.notion_prop_names.get('受文單位', '受文單位')
        while True:
            query_params = {"database_id": self.database_id, "page_size": 100};
            if use_filter: query_params["filter"] = {"property": archived_prop_name, "checkbox": {"equals": False}}
            if next_cursor: query_params["start_cursor"] = next_cursor
            try:
                response = self._notion_api_call(self.notion.databases.query, **query_params)
                results = response.get('results', []); total_retrieved += len(results); page_count += 1
                for page in results:
                    if page.get('archived', False): continue
                    properties = page.get('properties', {}); doc_id_prop = properties.get(title_prop_name); doc_type_prop = properties.get(doc_type_prop_name); recipient_prop = properties.get(recipient_prop_name)
                    doc_id, doc_type, recipient = None, None, None
                    if doc_id_prop and doc_id_prop.get('type') == 'title' and doc_id_prop.get('title'):
                        try: doc_id = doc_id_prop['title'][0].get('text', {}).get('content')
                        except: pass
                    if doc_type_prop and doc_type_prop.get('type') == 'select' and doc_type_prop.get('select'): doc_type = doc_type_prop['select'].get('name')
                    if recipient_prop:
                        prop_type = recipient_prop.get('type')
                        if prop_type == 'select' and recipient_prop.get('select'): recipient = recipient_prop['select'].get('name')
                        elif prop_type == 'rich_text' and recipient_prop.get('rich_text'):
                             try: recipient = recipient_prop['rich_text'][0].get('text', {}).get('content')
                             except: pass
                    recipient = str(recipient or '')
                    if doc_id and doc_type:
                        # --- 使用正規化 key (移除中間的'字') 來建立索引 ---
                        # 假設 Notion 中可能存在舊格式 "...第..." 或 新格式 "...字第..."
                        # 我們在生成 key 時，統一使用 "...第..." 格式
                        normalized_doc_id = doc_id.replace('字第', '第') if doc_id else None
                        if normalized_doc_id:
                             unique_key = f"{normalized_doc_id}_{doc_type}_{recipient}"
                             if unique_key in existing_docs: logger.warning(f"重複 unique_key (正規化後): {unique_key}. Page ID: {page.get('id')} 將覆蓋 {existing_docs[unique_key].get('id')}")
                             existing_docs[unique_key] = page # Use normalized key for index
                        else: logger.warning(f"頁面 {page.get('id', '未知ID')} 公文字號無法正規化: '{doc_id}'")
                    else: logger.warning(f"頁面 {page.get('id', '未知ID')} 缺少關鍵信息無法建立完整索引。")
                if response.get('has_more') and response.get('next_cursor'): next_cursor = response.get('next_cursor'); logger.debug(f"已獲取 {page_count} 頁 ({total_retrieved} 條)，繼續...")
                else: break
            except APIResponseError as e:
                error_body_str = str(e.body)
                if use_filter and archived_prop_name in error_body_str and ("Could not find property" in error_body_str or "is not a valid property" in error_body_str):
                    logger.error(f"獲取 Notion 文件失敗：資料庫缺少名為 '{archived_prop_name}' 的 Checkbox 屬性。將嘗試不使用過濾器...", exc_info=False)
                    use_filter = False; next_cursor = None; existing_docs = {}; page_count = 0; total_retrieved = 0; continue
                else: logger.error(f"Notion 獲取現有文件時發生 API 錯誤: Status={e.status}, Code={e.code}, Body={error_body_str}", exc_info=False); break
            except Exception as e: logger.error(f"Notion 獲取現有文件時出錯: {e}", exc_info=True); break
        logger.info(f"Notion 獲取完成: {page_count} 頁, {total_retrieved} 記錄 {'(含已歸檔)' if not use_filter else '(未歸檔)'}, 索引 {len(existing_docs)} 筆 (僅未歸檔)")
        return existing_docs

    def process_documents(self):
        """主處理流程"""
        start_time = time.time()
        try:
            logger.info("=== 開始公文同步流程 ===")
            date_files = self.get_date_files();
            if not date_files: logger.info("資料夾中無公文檔案。"); return
            all_processed_data_list = []
            logger.info("--- 階段 1: 讀取和準備本地 CSV 數據 ---")
            total_files = sum(len(files) for files in date_files.values()); processed_files = 0
            files_to_process = []
            for date_str, file_list in sorted(date_files.items()):
                for file_info in file_list: files_to_process.append((date_str, file_info))
            for date_str, file_info in tqdm(files_to_process, desc="處理 CSV 檔案", unit="file", leave=False):
                processed_files += 1; filename = os.path.basename(file_info["path"]); doc_type_detail = file_info["type"]
                try:
                    df = self.load_csv_data(file_info["path"], doc_type_detail)
                    if not df.empty:
                        prepared_df = self.prepare_data(df, doc_type_detail, date_str) # Generates "...字第..." format
                        if not prepared_df.empty: all_processed_data_list.append(prepared_df);
                        else: logger.warning(f"準備 {doc_type_detail} 數據失敗或結果為空: {filename}")
                except Exception as e: logger.error(f"處理 {doc_type_detail} 檔案時發生嚴重錯誤: {filename}, 錯誤: {e}", exc_info=True)

            if not all_processed_data_list: logger.warning("所有檔案處理後均無有效數據。"); return

            try:
                combined_raw_df = pd.concat(all_processed_data_list, ignore_index=True).fillna('')
                logger.info(f"數據合併完成，合併後總共 {len(combined_raw_df)} 筆記錄。")
                logger.info("開始進行跨日期重複數據處理 (保留最新記錄)...")
                original_count = len(combined_raw_df)
                combined_raw_df['系統輸出日期_dt'] = pd.to_datetime(combined_raw_df['系統輸出日期'], errors='coerce')
                # 使用原始公文字號排序，因為去重前需要原始格式
                combined_raw_df.sort_values(by=['公文字號', '文件類型', '受文單位', '系統輸出日期_dt'], ascending=[True, True, True, False], inplace=True, na_position='last')
                final_combined_df = combined_raw_df.drop(columns=['系統輸出日期_dt'])
                # --- 使用正規化後的公文字號進行去重 ---
                final_combined_df['normalized_doc_id'] = final_combined_df['公文字號'].str.replace('字第', '第', regex=False)
                final_combined_df = final_combined_df.drop_duplicates(subset=['normalized_doc_id', '文件類型', '受文單位'], keep='first').copy()
                final_combined_df.drop(columns=['normalized_doc_id'], inplace=True) # Remove temporary column
                # ------------------------------------
                deduplicated_count = len(final_combined_df)
                if original_count > deduplicated_count: logger.info(f"跨日期去重完成，保留 {deduplicated_count} 筆最新記錄。")
                else: logger.info(f"無需跨日期去重，待同步記錄數: {deduplicated_count}。")

                logger.info("正在生成流水號 (按公文日期排序)...")
                date_col_internal = '日期'
                final_combined_df[f'{date_col_internal}_dt'] = pd.to_datetime(final_combined_df[date_col_internal], errors='coerce')
                # 使用正規化後的公文字號進行排序以確保穩定性
                final_combined_df['normalized_doc_id_sort'] = final_combined_df['公文字號'].str.replace('字第', '第', regex=False)
                sort_keys_internal = [f'{date_col_internal}_dt', 'normalized_doc_id_sort', '文件類型', '受文單位']
                final_combined_df.sort_values(by=sort_keys_internal, ascending=[True, True, True, True], inplace=True, na_position='last')
                final_combined_df.reset_index(drop=True, inplace=True)
                final_combined_df['流水號'] = final_combined_df.index + 1
                final_combined_df.drop(columns=[f'{date_col_internal}_dt', 'normalized_doc_id_sort'], inplace=True)
                final_combined_df['流水號'] = final_combined_df['流水號'].astype(str)
                logger.info("流水號生成完畢。")

                doc_type_col_internal = '文件類型'
                invalid_rows = final_combined_df[~final_combined_df[doc_type_col_internal].isin(['發文', '收文'])];
                if not invalid_rows.empty: logger.error(f"!!! 去重後發現 '{doc_type_col_internal}' 列存在非預期值: {invalid_rows[doc_type_col_internal].unique()}，行索引: {invalid_rows.index.tolist()}")

            except Exception as final_proc_err: logger.critical(f"最終數據處理出錯: {final_proc_err}", exc_info=True); return

            logger.info("--- 階段 2: 連接 Notion 並獲取現有數據 ---")
            if not self.connect_notion(): return
            existing_docs = self.get_existing_documents() # Uses normalized keys
            logger.info(f"從 Notion獲取到 {len(existing_docs)} 筆現有(未歸檔)記錄。")

            logger.info("--- 階段 3: 同步數據到 Notion (創建/更新) ---")
            total_to_process = len(final_combined_df); created_count, updated_count, skipped_count, error_count = 0, 0, 0, 0
            processed_keys = set()
            failed_sync_info = []

            # --- !!! 修正點：修正 Create/Update 計數器邏輯 !!! ---
            for index, row in tqdm(final_combined_df.iterrows(), total=total_to_process, desc="同步 Notion", unit="rec", leave=False):
                doc_id_original = self.safe_get(row, '公文字號') # Local format (...字第...)
                doc_type = self.safe_get(row, '文件類型'); recipient = self.safe_get(row, '受文單位')
                # Use normalized key for lookup and tracking
                normalized_doc_id = doc_id_original.replace('字第', '第') if doc_id_original else None
                unique_key = f"{normalized_doc_id}_{doc_type}_{recipient}" if normalized_doc_id and doc_type else None
                progress_serial = row.get('流水號', index + 1);

                if doc_type not in ['發文', '收文']: logger.error(f"(流水號 {progress_serial}) 跳過: 無效文件類型 '{doc_type}'"); error_count += 1; failed_sync_info.append(f"流水號 {progress_serial}: 無效文件類型"); continue
                if not doc_id_original: logger.error(f"(流水號 {progress_serial}) 跳過: 公文字號為空"); error_count += 1; failed_sync_info.append(f"流水號 {progress_serial}: 公文字號為空"); continue
                if not unique_key: logger.error(f"(流水號 {progress_serial}) 跳過: 無法生成有效 Unique Key"); error_count += 1; failed_sync_info.append(f"流水號 {progress_serial}: 無法生成 Key"); continue

                processed_keys.add(unique_key) # Track processed normalized key
                logger.debug(f"處理: {unique_key} (流水號: {progress_serial})")
                try:
                    properties = self.build_notion_properties(row) # Build properties with "...字第..." format
                    if not properties: logger.error(f"跳過: 無法為 {unique_key} 構建有效屬性 ({progress_serial})。"); error_count += 1; failed_sync_info.append(f"{unique_key}: 屬性構建失敗"); continue

                    if unique_key in existing_docs: # Look up using normalized key
                        existing_page = existing_docs[unique_key]; page_id = existing_page['id']; logger.debug(f"已存在 (ID: {page_id})，檢查更新...")
                        if self.needs_update(existing_page, properties): # Compare against new properties (incl. new title format)
                            try:
                                update_payload = {}; existing_props = existing_page.get('properties', {})
                                title_prop_name_notion = self.notion_prop_names.get('公文字號', '公文字號')
                                for k_internal, v_dict in properties.items(): k_notion = self.notion_prop_names.get(k_internal, k_internal); update_payload[k_notion] = v_dict
                                # Check for properties to clear
                                new_prop_notion_names = set(update_payload.keys())
                                for k_notion, prop_data in existing_props.items():
                                    if k_notion != title_prop_name_notion and k_notion not in new_prop_notion_names:
                                        prop_type = prop_data.get('type'); has_value = False
                                        if prop_type == "select" and prop_data.get('select'): has_value = True
                                        elif prop_type == "date" and prop_data.get('date'): has_value = True
                                        elif prop_type == "number" and prop_data.get('number') is not None: has_value = True
                                        elif prop_type == "rich_text" and prop_data.get('rich_text'): has_value = True
                                        # Add other types...
                                        if has_value:
                                             clear_value = self._get_clear_payload(prop_type)
                                             if clear_value is not None or prop_type == 'select':
                                                  update_payload[k_notion] = clear_value; logger.debug(f"準備清空屬性 {k_notion} for page {page_id}")
                                # Execute Update
                                if self.update_notion_page(page_id, update_payload):
                                    updated_count += 1 # Increment correct counter
                                else: error_count += 1; failed_sync_info.append(f"{unique_key} (Update ID: {page_id}): API調用失敗")
                            except Exception as update_err: error_count += 1; failed_sync_info.append(f"{unique_key} (Update ID: {page_id}): {str(update_err)}")
                        else:
                            skipped_count += 1 # Increment correct counter
                    else:
                        logger.debug(f"不存在，嘗試創建...")
                        try:
                             new_page_id = self.create_notion_page(properties)
                             if new_page_id:
                                 created_count += 1 # Increment correct counter
                                 # Add to existing_docs using the *normalized* key so archive logic works if creation happens mid-run before a duplicate check
                                 existing_docs[unique_key] = {'id': new_page_id, 'properties': properties}
                             else: error_count += 1; failed_sync_info.append(f"{unique_key} (Create): API調用失敗或未返回ID")
                        except Exception as create_err: error_count += 1; failed_sync_info.append(f"{unique_key} (Create): {str(create_err)}")
                except Exception as row_proc_err: logger.error(f"處理 {unique_key} 時頂層錯誤: {row_proc_err}", exc_info=True); error_count += 1; failed_sync_info.append(f"{unique_key}: {str(row_proc_err)}"); continue
            # --------------------------------------------------

            # --- 階段 4 - 歸檔邏輯 ---
            logger.info("--- 階段 4: 歸檔 Notion 中不再存在於本地最新數據的記錄 ---")
            notion_keys = set(existing_docs.keys()) # Normalized keys from Notion
            # --- !!! 修正點：修正歸檔計數日誌 !!! ---
            keys_to_archive = notion_keys - processed_keys # Compare normalized keys
            archived_count = 0; archive_failed_count = 0
            logger.info(f"需要處理的本地最新記錄 Key 數量 (正規化): {len(processed_keys)}")
            logger.info(f"從 Notion 獲取的未歸檔記錄 Key 數量 (正規化): {len(notion_keys)}")
            logger.info(f"計算得出需要歸檔的 Notion 記錄數量: {len(keys_to_archive)}")
            # ---------------------------------------

            if keys_to_archive:
                 logger.info(f"開始執行歸檔 {len(keys_to_archive)} 筆記錄...")
                 for key_to_archive in tqdm(keys_to_archive, desc="執行歸檔", unit="page", leave=False):
                     page_to_archive = existing_docs.get(key_to_archive)
                     if page_to_archive and 'id' in page_to_archive:
                         page_id = page_to_archive['id']
                         logger.info(f"嘗試歸檔頁面: {key_to_archive} (ID: {page_id})")
                         try:
                              if self.archive_notion_page(page_id): archived_count += 1
                              else: archive_failed_count += 1; failed_sync_info.append(f"{key_to_archive} (Archive ID: {page_id}): API調用失敗")
                         except Exception as archive_err: archive_failed_count += 1; failed_sync_info.append(f"{key_to_archive} (Archive ID: {page_id}): {str(archive_err)}")
                     else: logger.warning(f"無法獲取需歸檔頁面 ID: {key_to_archive}"); archive_failed_count += 1; failed_sync_info.append(f"{key_to_archive} (Archive): 無法獲取 Page ID")
            else: logger.info("沒有需要歸檔的 Notion 記錄。")
            logger.info(f"歸檔操作完成: 成功歸檔 {archived_count} 筆, 歸檔失敗 {archive_failed_count} 筆。")

            # --- 階段 5 - 匯總 ---
            logger.info("--- 階段 5: 同步完成與結果匯總 ---")
            logger.info(f"總處理唯一記錄 (最新版本): {total_to_process} 筆")
            logger.info(f"新增到 Notion: {created_count} 筆")
            logger.info(f"更新 Notion: {updated_count} 筆")
            logger.info(f"無需更新(數據相同): {skipped_count} 筆")
            logger.info(f"歸檔 Notion 中舊記錄: {archived_count} 筆")
            logger.info(f"同步/歸檔失敗: {error_count + archive_failed_count} 筆")
            if failed_sync_info:
                logger.error("處理失敗的公文摘要 (最多顯示 10 條):")
                for i, failed_info in enumerate(failed_sync_info):
                    if i < 10: logger.error(f"- {failed_info}")
                    elif i == 10: logger.error(f"- ... (還有 {len(failed_sync_info) - 10} 條錯誤未顯示)") ; break

            # 6. 儲存 CSV (去重並按日期排序後的)
            try:
                output_file = os.path.join(self.folder_path, "combined_list_final.csv")
                final_columns_for_output = self.final_columns
                for col in final_columns_for_output:
                     if col not in final_combined_df.columns: logger.warning(f"最終 DataFrame 缺少欄位 '{col}'，添加空值輸出。"); final_combined_df[col] = ''
                output_df_to_save = final_combined_df[final_columns_for_output].fillna('')
                output_df_to_save.to_csv(output_file, index=False, encoding='utf-8-sig')
                logger.info(f"已儲存最終整合檔案 (包含 {len(output_df_to_save)} 筆唯一最新記錄，按日期排序): {output_file}")
            except Exception as save_err: logger.error(f"儲存最終整合檔案 '{output_file}' 失敗: {save_err}", exc_info=True)

            end_time = time.time(); logger.info(f"=== 公文同步流程結束 (耗時: {end_time - start_time:.2f} 秒) ===")
        except Exception as e: logger.critical(f"公文同步主流程未預期錯誤: {e}", exc_info=True); raise

    # --- Documentation Generation Method ---
    def generate_documentation(self, output_filename: str = "操作說明.docx"):
        """生成操作說明 Word 文件"""
        if not DOCX_AVAILABLE: logger.error("生成 Word 文件失敗：需要安裝 'python-docx' 庫 (pip install python-docx)"); print("錯誤：需要安裝 'python-docx'"); return False
        try:
            doc = docx.Document()
            def add_para(text, style=None, bold=False, italic=False): p = doc.add_paragraph(); run = p.add_run(text); p.style = style if style else None; run.bold = bold; run.italic = italic; return p
            def add_code_block(text): p = doc.add_paragraph(); run = p.add_run(text); run.font.name = 'Courier New'; run.font.size = Pt(10)
            doc.add_heading('公文同步 Notion 腳本操作說明', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading('1. 功能概述', level=2)
            add_para("本腳本旨在自動同步指定資料夾中的公文列表 CSV 檔案到指定的 Notion 資料庫。主要功能包括：")
            add_para("讀取多種類型的公文清單 CSV 檔案（電子收文/發文，紙本收文/發文）。", style='List Bullet')
            add_para("清理和轉換數據（如民國日期轉西元日期）。", style='List Bullet')
            add_para("處理跨越多個日期的重複公文記錄，只保留每個公文（以文號、類型、受文單位識別）最新日期的版本。", style='List Bullet')
            add_para("為最終的唯一記錄生成按公文日期排序的流水號。", style='List Bullet')
            add_para("將處理後的數據同步到 Notion 資料庫：創建新頁面、更新現有頁面、歸檔過時頁面。", style='List Bullet')
            add_para("將最終處理的唯一記錄列表保存為 combined_list_final.csv。", style='List Bullet')
            add_para("記錄詳細的執行日誌到 document_processing.log。", style='List Bullet')
            add_para("通過外部 config.json 文件進行配置。", style='List Bullet')
            add_para("在執行較長的操作時顯示進度條。", style='List Bullet')
            doc.add_heading('2. 環境準備', level=2)
            add_para('請確保系統已安裝 Python 3.8 或更高版本。')
            add_para('打開終端機，進入腳本所在資料夾，執行以下命令安裝必要的 Python 庫：')
            add_code_block("pip install pandas notion-client tenacity python-dotenv tqdm python-docx")
            doc.add_heading('3. 腳本與檔案設定', level=2)
            add_para('1.  腳本文件: 將 Officialdocument-Gemi.py 放置在你的工作目錄中。')
            add_para('2.  `.env` 環境變數檔:')
            add_para('    - 在與腳本相同的目錄下，創建 `.env` 文件。', style='List Bullet 2')
            add_para('    - 加入以下兩行並替換值：', style='List Bullet 2')
            add_code_block('NOTION_TOKEN="secret_YOUR_NOTION_TOKEN_HERE"\nNOTION_DATABASE_ID="YOUR_DATABASE_ID_HERE"')
            add_para('    - 前往 Notion Integrations 獲取 Token 並授權給目標 Database。', style='List Bullet 2')
            add_para('    - 在 Notion Database 頁面網址或分享鏈接中找到 Database ID。', style='List Bullet 2')
            add_para('3.  `config.json` 設定檔:')
            add_para('    - 在與腳本相同的目錄下，創建 `config.json` 檔案。', style='List Bullet 2')
            add_para('    - 複製以下範本內容，並根據說明修改：', style='List Bullet 2')
            config_to_show = getattr(self, 'config', DEFAULT_CONFIG)
            add_code_block(json.dumps(config_to_show, indent=2, ensure_ascii=False))
            add_para('    - **`folder_path`**: 修改為你的 CSV 資料夾路徑。', style='List Bullet 2')
            add_para('    - **`log_level`**: "INFO" 或 "DEBUG"。', style='List Bullet 2')
            add_para('    - **`notion_property_names`**: **重要！** 修改右側值為 Notion 中實際屬性名。**必須包含 "archived"**。', style='List Bullet 2')
            add_para('4.  數據資料夾與 CSV 檔案:')
            add_para('    - 確保 `folder_path` 指向的資料夾存在。', style='List Bullet 2')
            add_para('    - CSV 檔名格式: `YYYY-MM-DD_` + (`receiveList.csv`, `sendList.csv`, `preceiveList.csv`, 或 `psendList.csv`)。', style='List Bullet 2')
            add_para('    - CSV 內容從第 4 行開始為數據。', style='List Bullet 2')
            add_para('5.  Notion 資料庫設定:')
            add_para('    - 確保包含所有必要的屬性列，且名稱與 `config.json` 的值匹配，類型如下：', style='List Bullet 2')
            prop_list_str = "\n".join([f"        - {name}: {prop_type}" for name, prop_type in self.get_expected_notion_types().items()])
            add_para(prop_list_str)
            add_para('    - **重要**: 所有 Select 屬性，請預先添加選項。', style='List Bullet 2')
            doc.add_heading('4. 運行腳本', level=2)
            add_para('1.  打開終端機。')
            add_para('2.  `cd` 到腳本目錄。')
            add_para('3.  執行：'); add_code_block('python Officialdocument-Gemi.py')
            add_para('4.  觀察控制台輸出和進度條。')
            doc.add_heading('5. 輸出結果', level=2)
            add_para('- **Notion 資料庫**: 根據本地最新、唯一的數據進行更新。')
            add_para('- **`combined_list_final.csv`**: 在數據資料夾下生成/覆蓋，包含去重、排序後的唯一最新記錄。')
            add_para('- **`document_processing.log`**: 記錄詳細過程。')
            doc.add_heading('6. 故障排除', level=2)
            add_para('- **`SyntaxError`**: 檢查代碼複製是否完整。')
            add_para('- **`FileNotFoundError`**: 檢查 `config.json` 的 `folder_path` 和 `.env` 文件位置。')
            add_para('- **`ValueError`**: 檢查 `.env` 的 Token/Database ID 或 `config.json` 的格式。')
            add_para('- **`Notion API 錯誤`**: 查看日誌 `Body=` 內容。常見原因：權限、屬性不存在/類型錯誤、Select 選項未添加。**務必確認 `archived` Checkbox 屬性存在且名稱正確！**')
            add_para('- **歸檔不工作**: 確保 `archived` Checkbox 屬性存在。檢查日誌階段 4 輸出。')

            doc.save(output_filename)
            logger.info(f"操作說明文件已生成: {output_filename}")
            print(f"操作說明文件已生成: {output_filename}")
            return True
        except ImportError: logger.error("生成 Word 文件失敗：需要安裝 'python-docx' 庫 (pip install python-docx)"); print("錯誤：需要安裝 'python-docx'"); return False
        except Exception as e: logger.error(f"生成操作說明時發生錯誤: {e}", exc_info=True); print(f"生成操作說明時發生錯誤: {e}"); return False

    # --- Helper for Documentation ---
    def get_expected_notion_types(self) -> Dict[str, str]:
        """Returns a dictionary of expected Notion property names and types for documentation."""
        types = {}
        title_key = self.notion_prop_names.get('公文字號', '公文字號'); types[title_key] = "Title"
        select_keys = ["文件類型", "類別", "發文單位", "受文單位", "收發狀態", "發文形式", "承攬案件"]
        date_keys = ["日期", "收文日期", "系統輸出日期"]
        number_keys = ["流水號"]
        rich_text_keys = ["公文日期", "字", "文號", "主旨", "備註"]
        checkbox_keys = ["archived"]
        for key in select_keys: types[self.notion_prop_names.get(key, key)] = "Select"
        for key in date_keys: types[self.notion_prop_names.get(key, key)] = "Date"
        for key in number_keys: types[self.notion_prop_names.get(key, key)] = "Number"
        for key in rich_text_keys: types[self.notion_prop_names.get(key, key)] = "Rich Text / Text"
        for key in checkbox_keys: types[self.notion_prop_names.get(key, key)] = "Checkbox"
        return types

if __name__ == "__main__":
    print("開始執行公文同步腳本...")
    logger.info("腳本執行開始。")
    try:
        sync_processor = NotionDocumentSync(config)

        # --- 自動檢查並生成操作說明文檔 (如果不存在) ---
        doc_filename = "操作說明.docx"
        doc_path = os.path.join(sync_processor.folder_path, doc_filename)
        if not os.path.exists(doc_path):
            if DOCX_AVAILABLE:
                print(f"未找到操作說明文件 '{doc_path}'，正在生成...")
                logger.info(f"Attempting to generate documentation file at: {doc_path}")
                if sync_processor.generate_documentation(doc_path):
                     print(f"操作說明文件 '{doc_filename}' 已生成於 {sync_processor.folder_path}")
            else:
                 print("警告：未安裝 'python-docx' 庫，無法自動生成操作說明文件。請手動執行 'pip install python-docx'。")
                 logger.warning("python-docx library not found, skipping documentation generation.")
        else:
            print(f"操作說明文件 '{doc_path}' 已存在，跳過生成。")
            logger.info(f"Documentation file already exists at: {doc_path}")
        # ---------------------------------------------

        # --- 執行主要同步流程 ---
        sync_processor.process_documents()

    except ValueError as ve: logger.critical(f"腳本初始化失敗: {ve}"); print(f"腳本初始化失敗: {ve}")
    except FileNotFoundError as fnfe: logger.critical(f"腳本初始化失敗: {fnfe}"); print(f"腳本初始化失敗: {fnfe}")
    except Exception as main_exc: logger.critical(f"執行時發生嚴重錯誤: {main_exc}", exc_info=True); print(f"執行時發生嚴重錯誤，請檢查日誌。錯誤: {type(main_exc).__name__}")
    finally: logger.info("腳本執行完畢。"); print("腳本執行完畢。")